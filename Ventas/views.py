from django.shortcuts import render

# Create your views here.
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages

from Productos.models import Producto
from .models import Carrito, ItemCarrito


## Agregar al carrito REVISAR IA
@login_required
def agregar_al_carrito(request, producto_id):
    producto = get_object_or_404(Producto, id=producto_id)

    # ✅ Obtiene o crea un carrito temporal para el usuario
    carrito, _ = Carrito.objects.get_or_create(usuario=request.user)

    # ✅ Busca si el producto ya está en el carrito
    item, creado = ItemCarrito.objects.get_or_create(carrito=carrito, producto=producto)

    if not creado:
        # Si ya existe, aumenta la cantidad
        item.cantidad += 1
        item.save()
        messages.info(request, f"🔼 Se aumentó la cantidad de '{producto.nombre}' en tu carrito.")
    else:
        messages.success(request, f"🛒 '{producto.nombre}' se agregó al carrito.")

    # ✅ Redirige a la página desde donde vino el usuario (sin salir)
    return redirect(request.META.get('HTTP_REFERER', 'listar_categorias'))

#VER CARRITO REVISAR IA
@login_required
def ver_carrito(request):
    carrito, creado = Carrito.objects.get_or_create(usuario=request.user)
    items = carrito.items.select_related('producto')
    total = carrito.total()
    return render(request, 'ventas/ver_carrito.html', {
        'carrito': carrito,
        'items': items,
        'total': total
    })

#ACTUALIZAR CANTIDAD REVISAR IA
@login_required
def actualizar_cantidad(request, item_id, accion):
    item = get_object_or_404(ItemCarrito, id=item_id, carrito__usuario=request.user)

    if accion == 'sumar':
        item.cantidad += 1
    elif accion == 'restar':
        if item.cantidad > 1:
            item.cantidad -= 1
        else:
            item.delete()
            messages.info(request, "Producto eliminado del carrito.")
            return redirect('ver_carrito')

    item.save()
    return redirect('ver_carrito')

#@login_required
def eliminar_item(request, item_id):
    item = get_object_or_404(ItemCarrito, id=item_id, carrito__usuario=request.user)
    item.delete()
    messages.success(request, "Producto eliminado del carrito.")
    return redirect('ver_carrito')

#@login_required
def vaciar_carrito(request):
    carrito = get_object_or_404(Carrito, usuario=request.user)
    carrito.items.all().delete()
    messages.info(request, "Se vació tu carrito.")
    return redirect('ver_carrito')

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from io import BytesIO
import qrcode, base64
from Usuarios.models import Cliente, Empleado, CargoLaboral
from Ventas.models import Carrito, DetalleVenta, MetodoPago, Venta,Reserva, DetalleReserva, Notificacion
#from Notificaciones.models import Notificacion  # opcional si usas tabla Notificacion

@login_required
def pago_qr(request):
    carrito, _ = Carrito.objects.get_or_create(usuario=request.user)
    items = carrito.items.select_related('producto')
    total = carrito.total()

    cliente, creado = Cliente.objects.get_or_create(usuario=request.user)

    cargo_cajero = CargoLaboral.objects.filter(cargo__iexact="Cajero").first()
    cajeros = Empleado.objects.filter(estado=True, cargo=cargo_cajero)

    if request.method == "POST":
        id_cajero = request.POST.get("cajero")
        metodo = request.POST.get("metodo")

        # 🟦 GUARDAR EN SESIÓN
        request.session["cajero_id"] = id_cajero
        request.session["metodo_pago"] = metodo

        empleado = Empleado.objects.get(usuario_id=id_cajero)

        # ------------------------------------------------------
        # 1) RESERVA
        # ------------------------------------------------------
        if metodo == "reserva":
            reserva = Reserva.objects.create(
                cliente=cliente,
                cajero=empleado,
                total=total,
            )

            Notificacion.objects.create(
                cajero=empleado,
                mensaje=f"Nuevo pedido de reserva del cliente {request.user.username}."
            )

            for item in items:
                DetalleReserva.objects.create(
                    reserva=reserva,
                    producto=item.producto,
                    cantidad=item.cantidad
                )

            carrito.items.all().delete()

            return render(request, "ventas/reserva_confirmada.html", {
                "reserva": reserva,
                "cajero": empleado,
                "total": total
            })

        # ------------------------------------------------------
        # 2) PAYPAL
        # ------------------------------------------------------
        elif metodo == "paypal":
            return redirect("iniciar_pago_paypal")

        # ------------------------------------------------------
        # 3) QR BANCO UNIÓN
        # ------------------------------------------------------
        elif metodo == "qr":
            return redirect("pago_union")

    return render(request, "ventas/seleccionar_cajero.html", {
        "cajeros": cajeros,
        "total": total
    })
from django.utils.timezone import localdate


@login_required
def cajero_reservas(request):
    # Verificar que es empleado
    try:
        cajero = request.user.empleado
    except Empleado.DoesNotExist:
        messages.error(request, "No eres empleado.")
        return redirect("home")

    # Verificar que sea cajero
    if cajero.cargo.cargo != "Cajero":
        messages.error(request, "No eres cajero.")
        return redirect("home")

    # Fecha de hoy
    hoy = localdate()

    # Todas las reservas del cajero de hoy
    reservas_hoy = Reserva.objects.filter(
        cajero=cajero,
        fecha__date=hoy
    ).order_by("-fecha")

    # Separarlas por estado
    pendientes = reservas_hoy.filter(estado="Pendiente")
    confirmadas = reservas_hoy.filter(estado="Confirmada")

    return render(request, "ventas/cajero_reservas.html", {
        "pendientes": pendientes,
        "confirmadas": confirmadas,
    })

@login_required
def cajero_reserva_detalle(request, reserva_id):
    reserva = get_object_or_404(Reserva, id=reserva_id)

    try:
        cajero = request.user.empleado
    except Empleado.DoesNotExist:
        messages.error(request, "No eres empleado.")
        return redirect("home")

    # ✅ Seguridad: un cajero solo ve sus reservas
    if reserva.cajero != cajero:
        messages.error(request, "No tienes permiso para ver esta reserva.")
        return redirect("cajero_reservas")

    detalles = reserva.detalles.select_related("producto")

    return render(request, "ventas/cajero_reserva_detalle.html", {
        "reserva": reserva,
        "detalles": detalles
    })

@login_required
def marcar_entregado(request, venta_id):
    venta = get_object_or_404(Venta, id=venta_id)
    try:
        cajero = request.user.empleado
    except Empleado.DoesNotExist:
        messages.error(request, "No eres cajero.")
        return redirect("home")

    # autorización: solo cajero asignado o gerente puede marcar
    if venta.id_empleado != cajero and not request.user.es_gerente:
        messages.error(request, "No tienes permiso.")
        return redirect("cajero_reservas")

    venta.entregado = True
    venta.save()
    messages.success(request, f"Venta #{venta.id} marcada como entregada.")
    return redirect("cajero_reservas")

@login_required
def confirmar_reserva(request, reserva_id):
    reserva = get_object_or_404(Reserva, id=reserva_id)

    try:
        cajero = request.user.empleado
    except Empleado.DoesNotExist:
        messages.error(request, "No eres empleado.")
        return redirect("home")

    if reserva.cajero != cajero:
        messages.error(request, "No tienes permiso para confirmar esta reserva.")
        return redirect("cajero_reservas")

    if reserva.estado != "Pendiente":
        messages.warning(request, "La reserva ya fue procesada.")
        return redirect("cajero_reservas")

    # ✅ Crear venta
    metodo_pago, _ = MetodoPago.objects.get_or_create(descripcion="Reserva Confirmada")

    venta = Venta.objects.create(
        id_cliente=reserva.cliente,
        id_empleado=cajero,
        id_pago=metodo_pago,
        descuento=0,
        reserva=reserva  # 🔥 vincula la reserva con la venta
    )

    # ✅ Crear detalle de venta
    for det in reserva.detalles.all():
        DetalleVenta.objects.create(
            id_venta=venta,
            id_producto=det.producto,
            cantidad=det.cantidad
        )

    # ✅ Cambiar estado
    reserva.estado = "Confirmada"
    reserva.save()

    messages.success(request, f"Reserva #{reserva.id} confirmada correctamente.")
    return redirect("cajero_reservas")

@login_required
def mis_ventas(request):
    try:
        cliente = request.user.cliente
    except:
        messages.error(request, "No eres cliente.")
        return redirect("home")

    ventas = Venta.objects.filter(id_cliente=cliente).select_related('reserva').order_by('-fecha')

    return render(request, "ventas/mis_ventas.html", {"ventas": ventas})

@login_required
def detalle_venta(request, venta_id):
    venta = get_object_or_404(Venta, id=venta_id, id_cliente__usuario=request.user)
    detalles = venta.detalleventa_set.select_related('id_producto')
    return render(request, "ventas/detalle_venta.html", {
        "venta": venta,
        "detalles": detalles
    })

# ventas/views.py (fragmentos completos)
import paypalrestsdk
from decimal import Decimal
from django.shortcuts import redirect, render, get_object_or_404
from django.http import HttpResponse
from django.contrib import messages
from django.urls import reverse
from django.conf import settings
from django.contrib.auth.decorators import login_required

from Productos.models import Producto
from Usuarios.models import Cliente, Empleado
from .models import Carrito, ItemCarrito, Venta, DetalleVenta, MetodoPago, Reserva

from .paypal_utils import format_decimal_for_paypal

@login_required
def iniciar_pago_paypal(request):
    # CONFIGURAR PAYPAL
    paypalrestsdk.configure({
        "mode": settings.PAYPAL_MODE,
        "client_id": settings.PAYPAL_CLIENT_ID,
        "client_secret": settings.PAYPAL_CLIENT_SECRET,
    })
    """
    Inicia el pago por PayPal usando el carrito del usuario (Carrito, ItemCarrito).
    Método: POST recomendado (pero aquí también aceptamos GET para pruebas rápidas).
    """
    # Recuperar carrito del usuario (crea si no existe)
    carrito, _ = Carrito.objects.get_or_create(usuario=request.user)

    items_qs = carrito.items.select_related("producto").all()
    if not items_qs.exists():
        messages.error(request, "Tu carrito está vacío.")
        return redirect('ver_carrito')

    # Construir lista de items y calcular total
    items_for_paypal = []
    total = Decimal("0.00")

    for it in items_qs:
        prod = it.producto
        qty = Decimal(it.cantidad)
        price = Decimal(prod.precio_venta)
        subtotal = (price * qty).quantize(Decimal("0.01"))
        total += subtotal

        items_for_paypal.append({
            "name": prod.nombre,
            "sku": f"prod_{prod.id}",
            "price": format_decimal_for_paypal(price),
            "currency": "USD",  # Ver notas abajo sobre moneda
            "quantity": int(it.cantidad),
        })

    total_str = format_decimal_for_paypal(total)

    # Crear el objeto Payment de PayPal
    base_url = request.build_absolute_uri('/')[:-1]  # dominio con / al final; lo quitamos
    return_url = base_url + reverse('paypal_success')
    cancel_url = base_url + reverse('paypal_cancel')

    payment = paypalrestsdk.Payment({
        "intent": "sale",
        "payer": {"payment_method": "paypal"},
        "redirect_urls": {
            "return_url": return_url,   # PayPal redirige aquí tras aprobar el pago
            "cancel_url": cancel_url
        },
        "transactions": [{
            "item_list": {"items": items_for_paypal},
            "amount": {
                "total": total_str,
                "currency": "USD"
            },
            "description": f"Compra desde Supermercado - usuario: {request.user.username}"
        }]
    })

    if payment.create():
        # Guardar payment.id en session temporalmente para referencia (opcional)
        request.session['paypal_payment_id'] = payment.id
        # Encontrar link de aprobación y redireccionar
        for link in payment.links:
            if link.rel == "approval_url":
                approval_url = str(link.href)
                return redirect(approval_url)
        messages.error(request, "Error: no se encontró URL de aprobación en PayPal.")
        return redirect('ver_carrito')
    else:
        # Log o mostrar error
        messages.error(request, f"Error creando el pago PayPal: {payment.error}")
        return redirect('ver_carrito')


@login_required
def paypal_success(request):
    """
    PayPal redirect here after user approves the payment.
    Ejecuta el pago y crea Venta + DetalleVenta.
    """
    payment_id = request.GET.get('paymentId')
    payer_id = request.GET.get('PayerID')

    if not payment_id or not payer_id:
        messages.error(request, "Parámetros PayPal faltantes.")
        return redirect('ver_carrito')

    # Buscar el objeto Payment en PayPal
    payment = paypalrestsdk.Payment.find(payment_id)

    if not payment:
        messages.error(request, "No se encontró el pago en PayPal.")
        return redirect('ver_carrito')

    # Ejecutar el pago
    if payment.execute({"payer_id": payer_id}):
        # Pago exitoso -> crear Venta y DetalleVenta
        # Obtener carrito actual del usuario
        carrito = get_object_or_404(Carrito, usuario=request.user)
        items_qs = carrito.items.select_related("producto").all()
        if not items_qs.exists():
            messages.error(request, "Tu carrito ya está vacío (no se pudo crear la venta).")
            return redirect('listar_categorias')

        # Obtener/crear Cliente
        cliente, _ = Cliente.objects.get_or_create(usuario=request.user)

        # Seleccionar un empleado válido para asignar a la venta:
        # Intentamos un "Cajero" activo, si no existe usamos el primer empleado disponible.
        empleado = None
        try:
            empleado = Empleado.objects.filter(estado=True, cargo__cargo__iexact="Cajero").first()
            if not empleado:
                empleado = Empleado.objects.first()
        except Exception:
            empleado = None

        if empleado is None:
            # No podemos crear la venta si la FK es obligatoria; se puede ajustar el modelo en el futuro.
            messages.warning(request, "No hay empleados configurados. Crea al menos un Empleado para registrar la venta.")
            # Dejamos el carrito intacto para que el admin lo revise.
            return redirect('ver_carrito')

        metodo_pago, _ = MetodoPago.objects.get_or_create(descripcion="PayPal")

        # Crear la Venta
        venta = Venta.objects.create(
            id_cliente=cliente,
            id_empleado=empleado,
            id_pago=metodo_pago,
            descuento=0
        )

        # Crear detalles de venta
        for it in items_qs:
            DetalleVenta.objects.create(
                id_venta=venta,
                id_producto=it.producto,
                cantidad=it.cantidad
            )

        # Limpiar carrito
        carrito.items.all().delete()

        messages.success(request, f"Pago procesado correctamente. Venta #{venta.id} registrada.")
        # Redirige a detalle de venta o a historial
        return redirect('detalle_venta', venta_id=venta.id)
    else:
        messages.error(request, f"Error ejecutando el pago: {payment.error}")
        return redirect('ver_carrito')


@login_required
def paypal_cancel(request):
    messages.info(request, "Pago cancelado por el usuario.")
    return redirect('ver_carrito')

# ventas/views.py (añadir al final o en lugar adecuado)
import qrcode
import base64
from io import BytesIO
from django.http import HttpResponse, FileResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

from .emv_qr import build_emv_payload
from .models import Carrito, Venta, DetalleVenta, MetodoPago

@login_required
def pago_qr_interoperable(request):
    """
    Genera QR interoperable EMV y lo muestra; permite descargar PDF.
    Usa Carrito del usuario actual.
    """
    carrito, _ = Carrito.objects.get_or_create(usuario=request.user)
    items = carrito.items.select_related("producto")
    if not items.exists():
        messages.error(request, "Tu carrito está vacío.")
        return redirect("ver_carrito")

    total = carrito.total()
    # Datos del comercio - debes completar con los tuyos
    merchant_account = "77788899900"  # reemplaza por tu ID comercial o PSP
    merchant_name = "SupermercadoXYZ"
    merchant_city = "LA PAZ"

    payload = build_emv_payload(merchant_account, merchant_name, merchant_city, amount=total, currency="BOB")
    # Generar imagen QR
    qr = qrcode.make(payload)
    buffer = BytesIO()
    qr.save(buffer, format="PNG")
    b64 = base64.b64encode(buffer.getvalue()).decode()
    # render plantilla con qr_base64
    return render(request, "ventas/pago_qr_interoperable.html", {
        "qr_base64": b64,
        "total": total,
        "payload": payload,
        "merchant_name": merchant_name,
    })

@login_required
def descargar_pdf_qr(request):
    """
    Genera PDF con QR (a partir del payload que creamos). 
    Si quieres, puedes pasar merchant info o buscar del carrito.
    """
    carrito, _ = Carrito.objects.get_or_create(usuario=request.user)
    items = carrito.items.select_related("producto")
    if not items.exists():
        messages.error(request, "Tu carrito está vacío.")
        return redirect("ver_carrito")

    total = carrito.total()
    merchant_account = "77788899900"
    merchant_name = "SupermercadoXYZ"
    merchant_city = "LA PAZ"
    payload = build_emv_payload(merchant_account, merchant_name, merchant_city, amount=total, currency="BOB")

    # Generar QR image
    qr = qrcode.make(payload)
    buf = BytesIO()
    qr.save(buf, format="PNG")
    buf.seek(0)

    # Crear PDF en memoria
    pdf_buf = BytesIO()
    c = canvas.Canvas(pdf_buf, pagesize=A4)
    width, height = A4

    # Título
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 80, f"Pago QR - {merchant_name}")

    # Dibujar QR
    # reportlab necesita una imagen en disco o un BytesIO con ImageReader
    from reportlab.lib.utils import ImageReader
    img = ImageReader(buf)
    c.drawImage(img, 50, height - 350, width=250, height=250)

    # Texto del monto y payload
    c.setFont("Helvetica", 12)
    c.drawString(320, height - 150, f"Total: {total} BOB")
    c.drawString(320, height - 170, f"Ciudad: {merchant_city}")
    c.drawString(320, height - 190, f"Escanee con su app bancaria para pagar")

    c.showPage()
    c.save()
    pdf_buf.seek(0)

    filename = f"qr_pago_{request.user.username}.pdf"
    return FileResponse(pdf_buf, as_attachment=True, filename=filename)


@csrf_exempt
def webhook_pago_psp(request):
    """
    Endpoint para que tu PSP/banco notifique pagos.
    Debes dar esta URL al proveedor: /ventas/webhook/payments/
    El proveedor debe enviar JSON con al menos:
        { "venta_id": 123, "status": "PAID", "provider_ref": "ABC123" }
    """
    if request.method != "POST":
        return JsonResponse({"error": "method_not_allowed"}, status=405)

    try:
        data = json.loads(request.body)
    except Exception:
        return JsonResponse({"error": "invalid_json"}, status=400)

    # Validar: PSP puede enviar su firma, token, etc. IMPLEMENTA validación real en producción.
    venta_id = data.get("venta_id")
    status = data.get("status")
    provider_ref = data.get("provider_ref")

    if not venta_id:
        return JsonResponse({"error": "missing venta_id"}, status=400)

    try:
        venta = Venta.objects.get(id=venta_id)
    except Venta.DoesNotExist:
        return JsonResponse({"error": "venta_not_found"}, status=404)

    # Marcar pagado si status equivale a pago
    if status and status.upper() in ("PAID", "SUCCEEDED", "COMPLETED"):
        venta.estado_pago = "Pagado"
        venta.pago_referencia = provider_ref
        venta.save()
        return JsonResponse({"ok": True})

    return JsonResponse({"ok": False, "reason": "status_not_paid"})

@login_required
def pago_union(request):
    carrito, _ = Carrito.objects.get_or_create(usuario=request.user)
    items = carrito.items.select_related("producto")

    if not items.exists():
        messages.error(request, "Tu carrito está vacío.")
        return redirect("ver_carrito")

    total = sum(i.producto.precio_venta * i.cantidad for i in items)

    return render(request, "ventas/pago_union.html", {
        "total": total,
        "qr_path": "/media/qr_banco_union.jpeg"   # Tu QR estático real
    })

@login_required
def confirmar_pago_union(request):
    usuario = request.user

    # Recuperar cajero guardado
    cajero_id = request.session.get("cajero_id")

    if not cajero_id:
        messages.error(request, "No se seleccionó cajero para esta venta.")
        return redirect("ver_carrito")

    cajero = Empleado.objects.get(usuario__id=cajero_id)

    # Obtener carrito
    carrito = Carrito.objects.get(usuario=usuario)
    items = carrito.items.select_related("producto")

    cliente = Cliente.objects.get(usuario=usuario)

    metodo_pago, _ = MetodoPago.objects.get_or_create(
        descripcion="Pago QR Banco Unión"
    )

    # Crear venta con CAJERO
    venta = Venta.objects.create(
        id_cliente=cliente,
        id_empleado=cajero,      # ⬅ YA NO ES NULL
        id_pago=metodo_pago,
        estado_pago="Pendiente",
        entregado=False,
        pago_referencia="QR-BancoUnion"
    )

    # Crear detalles
    for item in items:
        DetalleVenta.objects.create(
            id_venta=venta,
            id_producto=item.producto,
            cantidad=item.cantidad
        )

    # Limpiar carrito
    carrito.items.all().delete()

    messages.success(request, f"Pago registrado. Venta #{venta.id} pendiente de verificación.")
    return redirect("mis_ventas")

@login_required
def ventas_cajero(request):
    try:
        cajero = request.user.empleado
    except:
        messages.error(request, "No eres empleado.")
        return redirect("home")

    ventas = Venta.objects.filter(id_empleado=cajero).order_by("-fecha")

    return render(request, "ventas/ventas_cajero.html", {
        "ventas": ventas
    })

from django.shortcuts import render, get_object_or_404
from .models import Venta, DetalleVenta

def todas_las_ventas(request):
    ventas = Venta.objects.all().order_by('-fecha')

    # 🔎 filtros
    nombre = request.GET.get("nombre", "").strip()
    fecha_desde = request.GET.get("desde", "")
    fecha_hasta = request.GET.get("hasta", "")

    # filtro por nombre de usuario del cliente
    if nombre:
        ventas = ventas.filter(id_cliente__usuario__username__icontains=nombre)

    # filtro por fecha DESDE
    if fecha_desde:
        ventas = ventas.filter(fecha__gte=fecha_desde)

    # filtro por fecha HASTA
    if fecha_hasta:
        ventas = ventas.filter(fecha__lte=fecha_hasta)

    return render(request, "ventas/todas_las_ventas.html", {
        "ventas": ventas,
        "filtro_nombre": nombre,
        "filtro_desde": fecha_desde,
        "filtro_hasta": fecha_hasta,
    })


def detalle_venta(request, venta_id):
    venta = get_object_or_404(Venta, id=venta_id)
    detalles = DetalleVenta.objects.filter(id_venta=venta)

    return render(request, "ventas/detalle_venta.html", {
        "venta": venta,
        "detalles": detalles
    })


def detalles_venta_ajax(request, venta_id):
    venta = get_object_or_404(Venta, id=venta_id)

    detalles = [{
        "producto": d.id_producto.nombre,
        "cantidad": float(d.cantidad),
        "precio": float(d.id_producto.precio_venta),
        "subtotal": float(d.cantidad * d.id_producto.precio_venta),
        "producto_id": d.id_producto.id,
    } for d in venta.detalleventa_set.all()]

    if hasattr(venta.fecha, "strftime"):
        fecha_str = venta.fecha.strftime("%Y-%m-%d")
    else:
        fecha_str = str(venta.fecha)[:10]

    return JsonResponse({
        "detalles": detalles,
        "fecha": fecha_str
    })
from django.http import JsonResponse

from django.views.decorators.csrf import csrf_exempt
from .models import Venta, DetalleVenta, Devolucion, DetalleDevolucion
from Productos.models import Producto
from Inventario.models import Retiro,DetalleRetiro
from django.db import transaction
from django.utils import timezone
from decimal import Decimal
@csrf_exempt
def devolver_producto(request, venta_id, producto_id):
    if request.method != "POST":
        return JsonResponse({"mensaje": "Método no permitido"}, status=405)

    try:
        venta = Venta.objects.get(id=venta_id)
        detalle = DetalleVenta.objects.get(id_venta=venta, id_producto_id=producto_id)
        producto = detalle.id_producto

        cantidad_devolver = float(request.POST.get("cantidad", 1))
        motivo = request.POST.get("motivo", "Sin motivo")
        accion = request.POST.get("accion")   # "inventario" o "retiro"

        # Validaciones
        if cantidad_devolver <= 0:
            return JsonResponse({"mensaje": "La cantidad debe ser mayor a 0"})

        if cantidad_devolver > detalle.cantidad:
            return JsonResponse({"mensaje": "No puedes devolver más de lo comprado"})

        with transaction.atomic():

            # Nueva devolución
            devolucion = Devolucion.objects.create()

            # Registrar detalle
            DetalleDevolucion.objects.create(
                id_devolucion=devolucion,
                id_detalle_venta=detalle,
                cantidad=cantidad_devolver,
                motivo=motivo
            )

            # Procesar inventario o retiro
            if accion == "inventario":
                producto.cantidad += Decimal(cantidad_devolver)
                producto.save()

            elif accion == "retiro":

                # Obtener cajero actual
                empleado = Empleado.objects.get(usuario=request.user)

                retiro = Retiro.objects.create(
                    tipo=Retiro.MANUAL,
                    empleado=empleado
                )

                DetalleRetiro.objects.create(
                    retiro=retiro,
                    producto=producto,
                    cantidad=cantidad_devolver,
                    motivo=motivo
                )

            # Actualizar detalle venta
            detalle.cantidad -= Decimal(cantidad_devolver)
            if detalle.cantidad <= 0:
                detalle.delete()
            else:
                detalle.save()

        return JsonResponse({"mensaje": "Devolución registrada correctamente"})


    except Exception as e:
        return JsonResponse({"mensaje": f"Error: {str(e)}"})
    
def lista_devoluciones(request):
    devoluciones = obtener_devoluciones_filtradas(request)

    # --- detectar exportaciones ---
    export = request.GET.get("export")

    if export == "excel":
        return exportar_devoluciones_excel(request)

    if export == "word":
        return exportar_devoluciones_word(request)

    if export == "pdf":
        return exportar_devoluciones_pdf(request)

    # --- si no exporta, solo mostrar la página ---
    return render(request, "ventas/devoluciones.html", {
        "devoluciones": devoluciones
    })

def obtener_devoluciones_filtradas(request):
    devoluciones = Devolucion.objects.all().order_by("-fecha")

    fecha = request.GET.get("fecha")
    desde = request.GET.get("desde")
    hasta = request.GET.get("hasta")
    producto = request.GET.get("producto")
    motivo = request.GET.get("motivo")

    if fecha:
        devoluciones = devoluciones.filter(fecha=fecha)

    if desde:
        devoluciones = devoluciones.filter(fecha__gte=desde)

    if hasta:
        devoluciones = devoluciones.filter(fecha__lte=hasta)

    if producto:
        devoluciones = devoluciones.filter(
            detalledevolucion__id_detalle_venta__id_producto__nombre__icontains=producto
        ).distinct()

    if motivo:
        devoluciones = devoluciones.filter(
            detalledevolucion__motivo__icontains=motivo
        ).distinct()

    return devoluciones

from openpyxl import Workbook
from django.http import HttpResponse
from Ventas.models import DetalleDevolucion

def exportar_excel(request):
    wb = Workbook()
    ws = wb.active
    ws.title = "Ventas"

    ws.append([
        "Venta ID", "Fecha", "Cliente", "Cajero", "Metodo de Pago",
        "Producto", "Cantidad", "Precio Venta", "Subtotal", "Descuento (%)", "Total Final"
    ])

    detalles = DetalleVenta.objects.select_related(
        "id_venta", "id_producto", "id_venta__id_cliente__usuario"
    )

    for det in detalles:
        venta = det.id_venta
        precio = float(det.id_producto.precio_venta)
        cantidad = float(det.cantidad)
        subtotal = precio * cantidad

        ws.append([
            venta.id,
            venta.fecha.strftime("%Y-%m-%d"),
            venta.id_cliente.usuario.username,
            venta.id_empleado.usuario.username,
            venta.id_pago.descripcion,
            det.id_producto.nombre,
            cantidad,
            precio,
            subtotal,
            float(venta.descuento),
            float(venta.total_venta)
        ])

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = 'attachment; filename="ventas.xlsx"'
    wb.save(response)
    return response

from docx import Document
from django.http import HttpResponse

def exportar_word(request):
    document = Document()
    document.add_heading("Reporte de Ventas", level=1)

    detalles = DetalleVenta.objects.select_related(
        "id_venta", "id_producto", "id_venta__id_cliente__usuario"
    )

    for det in detalles:
        venta = det.id_venta
        precio = float(det.id_producto.precio_venta)
        subtotal = float(det.cantidad) * precio

        p = document.add_paragraph()
        p.add_run(f"Venta #{venta.id}\n").bold = True
        p.add_run(f"Fecha: {venta.fecha}\n")
        p.add_run(f"Cliente: {venta.id_cliente.usuario.username}\n")
        p.add_run(f"Cajero: {venta.id_empleado.usuario.username}\n")
        p.add_run(f"Método de pago: {venta.id_pago.descripcion}\n")
        p.add_run(f"Producto: {det.id_producto.nombre}\n")
        p.add_run(f"Cantidad: {det.cantidad}\n")
        p.add_run(f"Precio: {precio}\n")
        p.add_run(f"Subtotal: {subtotal}\n")
        p.add_run(f"Descuento: {venta.descuento}%\n")
        p.add_run(f"Total final venta: {venta.total_venta}\n")
        document.add_paragraph("---------------------------------------------")

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response["Content-Disposition"] = 'attachment; filename=\"ventas.docx\"'
    document.save(response)
    return response

from reportlab.pdfgen import canvas
from django.http import HttpResponse

def exportar_pdf(request):
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="ventas.pdf"'

    p = canvas.Canvas(response)
    p.setFont("Helvetica", 12)

    y = 800
    p.drawString(50, y, "REPORTE DE VENTAS")
    y -= 40

    detalles = DetalleVenta.objects.select_related(
        "id_venta", "id_producto", "id_venta__id_cliente__usuario"
    )

    for det in detalles:
        venta = det.id_venta
        precio = float(det.id_producto.precio_venta)
        subtotal = float(det.cantidad) * precio

        if y < 100:
            p.showPage()
            p.setFont("Helvetica", 12)
            y = 800

        p.drawString(50, y, f"Venta #{venta.id}  |  Fecha: {venta.fecha}  |  Cliente: {venta.id_cliente.usuario.username}")
        y -= 20
        p.drawString(50, y, f"Cajero: {venta.id_empleado.usuario.username}  |  Pago: {venta.id_pago.descripcion}")
        y -= 20
        p.drawString(50, y, f"Producto: {det.id_producto.nombre}  | Cant: {det.cantidad} | Precio: {precio} | Subtotal: {subtotal}")
        y -= 20
        p.drawString(50, y, f"Descuento: {venta.descuento}% | Total Venta: {venta.total_venta}")
        y -= 40

    p.save()
    return response

def exportar_devoluciones_excel(request):
    devoluciones = obtener_devoluciones_filtradas(request)

    wb = Workbook()
    ws = wb.active
    ws.title = "Devoluciones"

    ws.append(["ID", "Fecha", "Producto", "Cantidad", "Motivo"])

    for dev in devoluciones:
        for det in dev.detalledevolucion_set.all():
            ws.append([
                dev.id,
                dev.fecha.strftime("%Y-%m-%d"),
                det.id_detalle_venta.id_producto.nombre,
                float(det.cantidad),
                det.motivo,
            ])

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = "attachment; filename=devoluciones.xlsx"

    wb.save(response)
    return response

def exportar_devoluciones_word(request):
    devoluciones = obtener_devoluciones_filtradas(request)

    document = Document()
    document.add_heading("Reporte de Devoluciones", level=1)

    for dev in devoluciones:
        document.add_heading(f"Devolución #{dev.id}", level=2)
        document.add_paragraph(f"Fecha: {dev.fecha}")

        for det in dev.detalledevolucion_set.all():
            p = document.add_paragraph()
            p.add_run(f"- Producto: {det.id_detalle_venta.id_producto.nombre}\n")
            p.add_run(f"  Cantidad: {det.cantidad}\n")
            p.add_run(f"  Motivo: {det.motivo}\n")

        document.add_paragraph("---------------------------------------------")

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="devoluciones.docx"'
    document.save(response)
    return response

def exportar_devoluciones_pdf(request):
    devoluciones = obtener_devoluciones_filtradas(request)

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = "attachment; filename=devoluciones.pdf"

    p = canvas.Canvas(response)
    p.setFont("Helvetica", 12)

    y = 800
    p.drawString(50, y, "REPORTE DE DEVOLUCIONES")
    y -= 40

    for dev in devoluciones:
        if y < 120:
            p.showPage()
            p.setFont("Helvetica", 12)
            y = 800

        p.drawString(50, y, f"Devolución #{dev.id} — {dev.fecha}")
        y -= 20

        for det in dev.detalledevolucion_set.all():
            p.drawString(
                70, y,
                f"- {det.id_detalle_venta.id_producto.nombre} | {det.cantidad} | {det.motivo}"
            )
            y -= 20

        y -= 15

    p.save()
    return response